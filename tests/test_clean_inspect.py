"""Layer-1 tests for `clean` and `inspect` (Slice 2).

Deterministic, offline, CI-safe. Builds a doc in-process, protects it with the
vanish vector, then asserts the Slice 2 contract:
  - inspect finds the injection (vector name, run id, text)
  - clean removes the injection (no markers, payload gone from extracted text)
  - clean restores structural equivalence: the cleaned doc's paragraphs equal the
    original's (the injected paragraph is fully removed, real content untouched)
  - clean on an unprotected doc is a harmless no-op
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from scarecrow.formats import docx as _docx  # noqa: F401  (registers vectors)
from scarecrow import vector
from scarecrow.core.marker import Marker, TAG_ATTR
from scarecrow.core.modes import Mode
from scarecrow.core.payloads import select_payload
from scarecrow.formats.docx.reverse import inspect_docx, clean_docx

ORIGINAL_PARAS = [
    "Board Memo",
    "Revenue was $4.2M this quarter.",
    "Second paragraph of real content.",
]


def _paras(path) -> list[str]:
    return [p.text for p in Document(path).paragraphs if p.text.strip()]


@pytest.fixture
def protected(tmp_path):
    """Write an original + a protected copy; return (original_path, protected_path)."""
    orig = tmp_path / "memo.docx"
    doc = Document()
    doc.add_heading(ORIGINAL_PARAS[0], 1)
    for p in ORIGINAL_PARAS[1:]:
        doc.add_paragraph(p)
    doc.save(str(orig))

    prot = tmp_path / "memo.scarecrow.docx"
    doc2 = Document(str(orig))
    v = vector.get("vanish")
    v.apply(doc2, select_payload(Mode.REFUSE, None), Marker())
    doc2.save(str(prot))
    return orig, prot


def test_inspect_finds_the_injection(protected):
    _, prot = protected
    found = inspect_docx(str(prot))
    assert len(found) == 1
    inj = found[0]
    assert inj.vector == "vanish"
    assert inj.run_id  # non-empty
    assert "AI processing refused" in inj.text


def test_inspect_reports_nothing_on_clean_doc(protected, tmp_path):
    orig, _ = protected
    assert inspect_docx(str(orig)) == []


def test_clean_removes_the_injection(protected, tmp_path):
    _, prot = protected
    out = tmp_path / "cleaned.docx"
    removed = clean_docx(str(prot), str(out))
    assert removed == 1
    # no markers left, and the payload is gone from the extracted text
    assert inspect_docx(str(out)) == []
    full = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "AI processing refused" not in full


def test_clean_restores_structural_equivalence(protected, tmp_path):
    orig, prot = protected
    out = tmp_path / "cleaned.docx"
    clean_docx(str(prot), str(out))
    # the injected paragraph is fully gone; the real content is identical
    assert _paras(str(out)) == _paras(str(orig)) == ORIGINAL_PARAS


def test_clean_on_unprotected_doc_is_noop(protected, tmp_path):
    orig, _ = protected
    out = tmp_path / "cleaned.docx"
    removed = clean_docx(str(orig), str(out))
    assert removed == 0
    assert _paras(str(out)) == ORIGINAL_PARAS
