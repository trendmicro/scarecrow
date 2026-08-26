"""Layer-1 test for the DOCX w:vanish vector (Slice 1).

Layer-1 = vector correctness, deterministic, runs in CI (no network, no LLM). For
the vanish vector this asserts the Slice 1 contract:
  - the payload appears in the extracted text stream (what an AI would ingest)
  - the human-visible body text is unchanged
  - the sc: marker is present on the injected run (so clean/inspect can find it)

The fixture is built in-process (python-docx) rather than committed as a binary,
so the test is transparent and diffable.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

# Importing the format package registers the vector.
from scarecrow.formats import docx as _docx  # noqa: F401
from scarecrow import vector
from scarecrow.core.marker import Marker, TAG_ATTR, VECTOR_ATTR
from scarecrow.core.modes import Mode
from scarecrow.core.payloads import select_payload

VISIBLE = "Q3 revenue was $4.2M, up 18% quarter over quarter."


@pytest.fixture
def protected_doc():
    """A one-paragraph doc, protected with the vanish vector; returns (doc, canary)."""
    doc = Document()
    doc.add_paragraph(VISIBLE)
    marker = Marker()
    payload = select_payload(Mode.REFUSE, None)
    vector.get("vanish").apply(doc, payload, marker)
    # Round-trip through save/load so we test what an extractor actually sees.
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf), marker.run_id


def _extracted_text(doc) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def test_payload_is_in_extracted_text(protected_doc):
    doc, _ = protected_doc
    text = _extracted_text(doc)
    assert "PROTECTED DOCUMENT: AI processing refused" in text


def test_visible_content_unchanged(protected_doc):
    doc, _ = protected_doc
    text = _extracted_text(doc)
    assert VISIBLE in text


def test_injected_run_is_hidden(protected_doc):
    """The payload run must carry w:vanish (invisible in Word)."""
    doc, _ = protected_doc
    hidden_runs = [r for p in doc.paragraphs for r in p.runs if r.font.hidden]
    assert hidden_runs, "expected at least one hidden (w:vanish) run"
    assert any("processing refused" in r.text for r in hidden_runs)


def test_marker_present_on_injected_run(protected_doc):
    """The sc: marker (tag + vector name) must be on the injected run."""
    doc, run_id = protected_doc
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tagged = [
        r for p in doc.paragraphs for r in p.runs
        if r._r.get(TAG_ATTR) == run_id and r._r.get(VECTOR_ATTR) == "vanish"
    ]
    assert len(tagged) >= 1


def test_vanish_vector_is_registered():
    names = [v.name for v in vector.all_vectors("docx")]
    assert "vanish" in names
