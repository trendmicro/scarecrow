"""Layer-1 tests for the DOCX vector library + selection + sidecar (Slice 3).

Deterministic, offline. Covers:
  - all six DOCX vectors are registered
  - each vector applied alone: its payload reaches the file, and clean removes it
    (structural equivalence restored; sidecar restored for the text-value vector)
  - --all applies the default vector set and clean reverses all of them
  - the sidecar file is written iff a text-value vector ran
  - vector selection resolves names / all / default, and rejects bad input
"""
from __future__ import annotations

import zipfile

import pytest
from docx import Document

from scarecrow.formats import docx as _docx  # noqa: F401  (registers vectors)
from scarecrow import vector
from scarecrow.core.marker import Marker, write_sidecar, sidecar_path_for
from scarecrow.core.modes import Mode
from scarecrow.core.payloads import select_payload
from scarecrow.formats.docx.reverse import inspect_docx, clean_docx

ALL = ["vanish", "white_text", "header_footer", "off_page", "alt_text", "core_props",
       "unicode_tags", "zero_width", "comments", "tracked_delete"]
# Vectors whose payload is encoded (not literal) in the file — checked differently.
ENCODED = {"unicode_tags", "zero_width"}
PAYLOAD_MARK = "protected against automated AI"  # appears in the full refuse payload
ORIGINAL = ["Report", "Real line one.", "Real line two."]


def _make_original(tmp_path):
    p = tmp_path / "doc.docx"
    d = Document()
    d.add_heading(ORIGINAL[0], 1)
    for line in ORIGINAL[1:]:
        d.add_paragraph(line)
    d.save(str(p))
    return p


def _protect(orig, tmp_path, vectors: list[str]):
    """Apply the named vectors to a copy; write sidecar; return the protected path."""
    out = tmp_path / "doc.scarecrow.docx"
    doc = Document(str(orig))
    marker = Marker()
    payload = select_payload(Mode.REFUSE, None)
    for name in vectors:
        vector.get(name, "docx").apply(doc, payload, marker)
    doc.save(str(out))
    write_sidecar(str(out), marker.sidecar_entries)
    return out


def _all_xml(path) -> str:
    z = zipfile.ZipFile(str(path))
    return "".join(z.read(n).decode("utf-8", "replace")
                   for n in z.namelist() if n.endswith(".xml"))


def _paras(path):
    return [p.text for p in Document(str(path)).paragraphs if p.text.strip()]


def test_all_vectors_registered():
    names = [v.name for v in vector.all_vectors("docx")]
    assert names == ALL  # order matters for --all / list-vectors


def _decode_present(xml: str) -> bool:
    """True if an encoded payload (Tags-block or zero-width) is present in xml."""
    from_tags = "".join(
        chr(ord(c) - 0xE0000) if 0xE0000 <= ord(c) <= 0xE007F else ""
        for c in xml)
    if "protected" in from_tags.lower():
        return True
    # zero-width: just confirm a meaningful run of ZW chars exists
    return (xml.count("​") + xml.count("‌")) > 100


@pytest.mark.parametrize("name", ALL)
def test_vector_payload_reaches_the_file(name, tmp_path):
    orig = _make_original(tmp_path)
    prot = _protect(orig, tmp_path, [name])
    xml = _all_xml(prot)
    if name in ENCODED:
        assert _decode_present(xml)  # payload is encoded, not literal
    else:
        # core_props uses a shortened message; every vector still lands *some* notice
        assert ("protected" in xml.lower()) or ("owner-protected" in xml.lower())


@pytest.mark.parametrize("name", ALL)
def test_vector_clean_restores_original(name, tmp_path):
    orig = _make_original(tmp_path)
    prot = _protect(orig, tmp_path, [name])
    cleaned = tmp_path / "cleaned.docx"
    clean_docx(str(prot), str(cleaned))
    # structural equivalence + no payload left anywhere
    assert _paras(cleaned) == ORIGINAL
    xml = _all_xml(cleaned).lower()
    assert "protected against automated ai" not in xml
    assert "owner-protected" not in xml
    if name in ENCODED:
        assert not _decode_present(_all_xml(cleaned))  # encoded payload gone too


def test_all_vectors_apply_and_clean_reverses(tmp_path):
    # every vector (including the research-only ones) can be applied and cleaned
    orig = _make_original(tmp_path)
    prot = _protect(orig, tmp_path, ALL)
    found = inspect_docx(str(prot))
    assert {i.vector for i in found} == set(ALL)
    cleaned = tmp_path / "cleaned.docx"
    removed = clean_docx(str(prot), str(cleaned))
    assert removed == len(ALL)
    assert inspect_docx(str(cleaned)) == []
    assert _paras(cleaned) == ORIGINAL


def test_default_set_excludes_research_only_encoding_vectors():
    # --all applies the default posture, which drops the two encoding vectors
    # (no model decoded them). They remain in the full library.
    default_names = [v.name for v in vector.default_vectors("docx")]
    assert set(ENCODED).isdisjoint(default_names)          # excluded
    assert set(default_names) == set(ALL) - ENCODED        # everything else kept
    assert set(ENCODED).issubset(v.name for v in vector.all_vectors("docx"))  # still selectable


def test_sidecar_written_only_for_text_value_vectors(tmp_path):
    orig = _make_original(tmp_path)
    # structural-only: no sidecar
    prot1 = _protect(orig, tmp_path, ["vanish"])
    assert not sidecar_path_for(str(prot1)).exists()
    # includes the text-value core_props: sidecar exists with one entry
    prot2 = _protect(orig, tmp_path, ["vanish", "core_props"])
    sp = sidecar_path_for(str(prot2))
    assert sp.exists()


def test_inspect_reports_vector_and_location(tmp_path):
    orig = _make_original(tmp_path)
    prot = _protect(orig, tmp_path, ["header_footer", "core_props"])
    found = {i.vector: i.where for i in inspect_docx(str(prot))}
    assert found["header_footer"] == "header"
    assert found["core_props"].startswith("sidecar:")
